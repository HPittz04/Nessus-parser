import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import xml.etree.ElementTree as ET

# Função para definir a cor de fundo de uma célula usando XML
def set_cell_background_color(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_shading.set(qn('w:val'), 'clear')
    cell_properties.append(cell_shading)

# Função para definir a cor do texto de uma célula
def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color

# Função para substituir placeholders dentro de parágrafos e células de tabela
def replace_placeholders(paragraph, replacements):
    for key, value in replacements.items():
        if f'{{{{{key}}}}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', value)

# Função para justificar o texto à esquerda
def set_paragraph_left_aligned(paragraph):
    """
    Define o alinhamento do parágrafo para a esquerda.
    Args:
        paragraph: Objeto de parágrafo (docx.paragraph).
    """
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Função para extrair informações do arquivo .nessus
def extract_nessus_data(file_path):
    """
    Extrai as informações de um arquivo .nessus (XML) e retorna um dicionário com os dados.
    Args:
        file_path: Caminho para o arquivo .nessus.
    Returns:
        Lista de dicionários com informações da vulnerabilidade.
    """
    vulnerabilities = []
    
    # Parse do arquivo .nessus como XML
    tree = ET.parse(file_path)
    root = tree.getroot()

    # Loop através dos elementos 'ReportItem' para extrair as informações
    for report_item in root.iter('ReportItem'):
        severity_code = report_item.attrib.get('severity', '0')
        severity_text = get_severity_level(severity_code)
        
        # Filtra apenas vulnerabilidades com severidade Medium, High ou Critical
        if severity_text in ['Medium', 'High', 'Critical']:
            vulnerability = {
                "NOME_VULNERABILIDADE": report_item.attrib.get('pluginName', 'Unknown'),   # Nome do plugin
                "SEVERIDADE": severity_text,  # Severidade (convertida para texto)
                "CATEGORIA": report_item.attrib.get('pluginFamily', 'Unknown'),            # Família/Categoria do plugin
                "DESCRICAO": report_item.find('description').text if report_item.find('description') is not None else '',
                "RECOMENDACOES": report_item.find('solution').text if report_item.find('solution') is not None else ''
            }
            vulnerabilities.append(vulnerability)

    return vulnerabilities

# Função auxiliar para converter severidade de número para texto
def get_severity_level(severity_code):
    severity_levels = {
        '0': 'Informational',
        '1': 'Low',
        '2': 'Medium',
        '3': 'High',
        '4': 'Critical'
    }
    return severity_levels.get(severity_code, 'Unknown')

# Mapeamento de cores para severidade
severity_color_map = {
    "Medium": "F79646",  # Laranja
    "High": "DD4B50",    # Vermelho
    "Critical": "91243E" # Vermelho Escuro
}

# Função principal para preencher o template a partir de um arquivo .nessus
def fill_template_from_nessus(template_path, nessus_file, output_path=None):
    # Carrega o documento de template
    doc = Document(template_path)
    
    # Extrai os dados do arquivo .nessus com filtro de severidade
    vulnerabilities = extract_nessus_data(nessus_file)
    
    # Ordenar as vulnerabilidades por severidade: Critical -> High -> Medium
    vulnerabilities.sort(key=lambda x: {'Critical': 3, 'High': 2, 'Medium': 1}[x["SEVERIDADE"]], reverse=True)

    # Contar quantas tabelas existem no template e quantas vulnerabilidades foram filtradas
    num_tables = len(doc.tables)
    num_vulnerabilities = len(vulnerabilities)
    print(f"Total de tabelas no template: {num_tables}")
    print(f"Total de vulnerabilidades filtradas (Medium ou superior): {num_vulnerabilities}")

    # Preencher o template com as informações de cada vulnerabilidade
    vuln_index = 0  # Índice para acompanhar qual vulnerabilidade estamos preenchendo

    # Iterar sobre as tabelas e preencher cada vulnerabilidade
    for table in doc.tables:
        # Verificar se ainda há vulnerabilidades a serem preenchidas
        if vuln_index < num_vulnerabilities:
            vulnerability = vulnerabilities[vuln_index]  # Seleciona a vulnerabilidade atual
            
            # Preencher a tabela com os dados da vulnerabilidade
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Substituir os placeholders e aplicar o alinhamento à esquerda
                        replace_placeholders(paragraph, vulnerability)
                        set_paragraph_left_aligned(paragraph)

                    # Aplicar cor de fundo e texto se for uma célula de severidade
                    if '{{SEVERIDADE}}' in cell.text or cell.text == vulnerability["SEVERIDADE"]:
                        severity = vulnerability["SEVERIDADE"]
                        if severity in severity_color_map:
                            set_cell_background_color(cell, severity_color_map[severity])

                            # Alterar cor do texto para branco se a severidade for Critical
                            if severity == "Critical":
                                set_cell_text_color(cell, RGBColor(255, 255, 255))

            vuln_index += 1  # Avançar para a próxima vulnerabilidade

        # Se todas as vulnerabilidades foram processadas, sair do loop
        if vuln_index >= num_vulnerabilities:
            break

    # Verificar se alguma vulnerabilidade foi omitida por falta de tabelas no template
    if vuln_index < num_vulnerabilities:
        print(f"⚠️ Atenção: O template possui apenas {num_tables} tabelas, mas há {num_vulnerabilities} vulnerabilidades. {num_vulnerabilities - vuln_index} vulnerabilidades não foram preenchidas no relatório.")

    # Gerar o nome do arquivo de saída baseado no nome do arquivo .nessus, se não for fornecido
    if output_path is None:
        base_name = os.path.splitext(os.path.basename(nessus_file))[0]  # Extrair nome base do arquivo .nessus
        output_path = f"{base_name}-report.docx"  # Adicionar sufixo "-report"

    # Salvar o documento preenchido e atualizado
    doc.save(output_path)
    #print(f"Documento preenchido e salvo como '{output_path}'")
    print(f"Document saved as '{output_path}'")

# Executar a função principal com os caminhos especificados
template_file = 'template.docx'
nessus_file = 'file.nessus'
fill_template_from_nessus(template_file, nessus_file)
