import xml.etree.ElementTree as ET
import os
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from lxml import etree

def extrair_vulnerabilidades_e_ips(arquivo_nessus):
    try:
        # Carregar o arquivo XML
        tree = ET.parse(arquivo_nessus)
        root = tree.getroot()

        # Lista para armazenar vulnerabilidades e seus detalhes
        dados = []

        # Navegar pela árvore XML
        for report_host in root.findall(".//ReportHost"):
            host_name = report_host.get('name')  # Nome do host

            # Iterar sobre as vulnerabilidades
            for report_item in report_host.findall(".//ReportItem"):
                # Extrair informações relevantes da vulnerabilidade
                plugin_id = report_item.get('pluginID')
                plugin_name = report_item.get('pluginName')
                severity = report_item.get('severity')
                publication_date = report_item.find('.//plugin_publication_date').text if report_item.find('.//plugin_publication_date') is not None else "Data não disponível"


                # Filtrar apenas vulnerabilidades com severidade High ou maior
                if severity in ['2', '3', '4']:  # 2 = Medium, 3 = High, 4 = Critical
                    severidade_texto = "Medium" if severity == '2' else "High" if severity == '3' else "Critical"
                    dados.append({
                        'severity': severidade_texto,
                        'plugin_id': plugin_id,
                        'plugin_name': plugin_name,
                        'hosts': host_name,
                        'publication_date': publication_date
                    })

        return dados

    except ET.ParseError as e:
        print(f"Erro ao analisar o arquivo XML: {e}")
    except Exception as e:
        print(f"Ocorreu um erro: {e}")

def set_cell_background_color(cell, color):
    """Define a cor de fundo da célula usando lxml."""
    cell_element = cell._element
    tc_pr = cell_element.get_or_add_tcPr()

    # Criar o elemento de sombra (shd) e definir a cor
    shd = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd", 
                        {'fill': color})
    
    # Adicionar o elemento shd ao tcPr
    tc_pr.append(shd)

def salvar_em_word(dados, arquivo_saida):
    # Criar um novo documento Word
    doc = Document()
    #doc.add_heading('Relatório de Vulnerabilidades', level=1)
    doc.add_heading('Vulnerability Report', level=1)

    # Adicionar uma tabela
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'  # Estilo da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity (CVSS v3.0)'
    hdr_cells[1].text = 'Plugin ID'
    hdr_cells[2].text = 'Plugin Name'
    hdr_cells[3].text = 'Hosts'
    hdr_cells[4].text = 'Count'
    hdr_cells[5].text = 'Publication Date'

    # Dicionário para contar hosts por vulnerabilidade
    count_hosts = {}

    for item in dados:
        plugin_key = (item['plugin_id'], item['plugin_name'], item['severity'], item['publication_date'])
        if plugin_key not in count_hosts:
            count_hosts[plugin_key] = []
        count_hosts[plugin_key].append(item['hosts'])

    # Adicionar dados à tabela
    for (plugin_id, plugin_name, severity, publication_date), hosts in count_hosts.items():
        row_cells = table.add_row().cells

        # Adicionar severidade com cor de fundo
        row_cells[0].text = severity
        if severity == "Critical":
            fill_color = "91243E"  # Vermelho Escuro
        elif severity == "High":
            fill_color = "DD4B50"  # Vermelho
        elif severity == "Medium":
            fill_color = "F79646"  # Laranja
        else:
            fill_color = "FFFFFF"  # Branco por padrão

        # Definindo a cor de fundo da célula
        set_cell_background_color(row_cells[0], fill_color)
        
        # Preencher os outros campos
        row_cells[1].text = plugin_id
        row_cells[2].text = plugin_name
        row_cells[3].text = ', '.join(hosts)
        row_cells[4].text = str(len(hosts))
        row_cells[5].text = publication_date
        
        # Alinhamento vertical central
        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Salvar o documento
    doc.save(arquivo_saida)

# Caminho para o arquivo .nessus
arquivo_nessus = 'file.nessus'
# Caminho para o arquivo de saída
arquivo_saida = 'Vulnerability_Report.docx'

# Verificar se o arquivo existe
if os.path.exists(arquivo_nessus):
    dados_encontrados = extrair_vulnerabilidades_e_ips(arquivo_nessus)
    salvar_em_word(dados_encontrados, arquivo_saida)
    #print(f"Relatório salvo em: {arquivo_saida}")
    print(f"Report saved in: {arquivo_saida}")
else:
    #print("Arquivo não encontrado:", arquivo_nessus)
    print("File not found:", arquivo_nessus)
